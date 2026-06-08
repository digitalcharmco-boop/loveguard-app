#!/usr/bin/env python3
"""
YouTube Channel Clone Engine — Execution Layer
Deterministic AI operations for channel analysis, script generation,
visual profiling, and content production.
"""

import json
import os
import re
import time
import logging
from io import BytesIO
from typing import Dict, List, Optional

import openai
from dotenv import load_dotenv

load_dotenv()

logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)


class YouTubeCloneEngine:
    def __init__(self, api_key: str = None):
        self.api_key = api_key or os.getenv("OPENAI_API_KEY")
        if not self.api_key:
            raise ValueError("OPENAI_API_KEY environment variable not set")
        self.client = openai.OpenAI(api_key=self.api_key)
        self.model = "gpt-4o"
        self.max_retries = 3
        self.retry_delay = 1

    # ── Internal helpers ────────────────────────────────────────────────────

    def _call_api(
        self,
        messages: List[Dict],
        temperature: float = 0.7,
        max_tokens: int = 4096,
    ) -> str:
        for attempt in range(self.max_retries):
            try:
                response = self.client.chat.completions.create(
                    model=self.model,
                    messages=messages,
                    temperature=temperature,
                    max_tokens=max_tokens,
                )
                return response.choices[0].message.content
            except openai.RateLimitError:
                if attempt < self.max_retries - 1:
                    time.sleep(self.retry_delay * (2**attempt))
                else:
                    raise
            except openai.APIError as e:
                logger.error(f"OpenAI API error (attempt {attempt + 1}): {e}")
                if attempt < self.max_retries - 1:
                    time.sleep(self.retry_delay)
                else:
                    raise
        return ""

    def _parse_json(self, text: str):
        """Strip markdown fences and parse JSON."""
        text = text.strip()
        text = re.sub(r"^```(?:json)?\s*", "", text)
        text = re.sub(r"\s*```$", "", text)
        return json.loads(text)

    # ── State 4: Style DNA Extraction ────────────────────────────────────

    def extract_style_dna(self, transcripts: List[str], channel_url: str = "") -> Dict:
        non_empty = [t for t in transcripts if t and t.strip()]
        if non_empty:
            combined = "\n\n--- TRANSCRIPT SEPARATOR ---\n\n".join(
                f"TRANSCRIPT {i + 1}:\n{t}" for i, t in enumerate(non_empty)
            )
            content_source = f"Channel transcripts/descriptions:\n\n{combined}"
        else:
            content_source = (
                f"No transcripts available. Infer style DNA from the channel URL alone: {channel_url}\n"
                "Use your knowledge of this channel's known style, pacing, tone, and format."
            )
        messages = [
            {
                "role": "system",
                "content": (
                    "You are an elite YouTube content strategist. Analyze transcripts and "
                    "extract style DNA. Deconstruct mechanics, not content. Output pure JSON only."
                ),
            },
            {
                "role": "user",
                "content": f"""Analyze this YouTube channel.
Channel: {channel_url}

{content_source}

Extract Style DNA as JSON with EXACTLY these keys:
{{
  "niche": "Core subject territory",
  "target_audience": "Who is spoken to and how",
  "hook_architecture": "How attention is captured in the first 15-30 seconds",
  "script_flow": "Structural pattern from open to close",
  "sentence_rhythm": "Length variation, cadence, pacing patterns",
  "tone_profile": "Formal/casual, authoritative/curious, warm/urgent",
  "transition_techniques": "How the creator moves between ideas",
  "curiosity_gaps": "How information is withheld or teased",
  "emotional_triggers": "Fear/aspiration/surprise/validation — where and how",
  "retention_devices": "Loops, callbacks, pattern interrupts, re-engagement cues",
  "direct_address_style": "How the audience is spoken to personally",
  "words_per_second": 2.2,
  "target_word_count": 1200,
  "overall_voice": "One-paragraph synthesis of the channel's unique voice"
}}

Set words_per_second and target_word_count as numbers derived from the transcripts.
Output only JSON.""",
            },
        ]
        result = self._call_api(messages, temperature=0.3)
        return self._parse_json(result)

    # ── State 3 helper: Topic Ideas ──────────────────────────────────────

    def generate_topic_ideas(self, style_dna: Dict) -> List[str]:
        messages = [
            {
                "role": "system",
                "content": "You are an elite YouTube content strategist. Output pure JSON only.",
            },
            {
                "role": "user",
                "content": f"""Based on this channel's Style DNA, generate 5 compelling video topic ideas.

Style DNA:
{json.dumps(style_dna, indent=2)}

Output a JSON array of 5 strings. Each string is a specific, compelling video title/topic.
Output only the JSON array.""",
            },
        ]
        result = self._call_api(messages, temperature=0.8)
        return self._parse_json(result)

    # ── State 5: Script Generation ───────────────────────────────────────

    def generate_script(self, style_dna: Dict, topic: str) -> Dict:
        target_wc = int(style_dna.get("target_word_count", 1200))
        low, high = int(target_wc * 0.95), int(target_wc * 1.05)

        messages = [
            {
                "role": "system",
                "content": (
                    "You are an elite YouTube scriptwriter. Write complete, publish-ready scripts "
                    "that replicate a channel's DNA. Output script text only — no labels, "
                    "no headings, no commentary, no visual direction."
                ),
            },
            {
                "role": "user",
                "content": f"""Write a complete YouTube script for this topic: "{topic}"

Style DNA to replicate:
{json.dumps(style_dna, indent=2)}

ABSOLUTE RULES:
- Target: {target_wc} words (acceptable range: {low}–{high})
- Mirror hook architecture, script flow, sentence rhythm, tone, and emotional arc exactly
- Use the same transition techniques and retention devices
- Zero generic YouTube templates
- Zero visual direction or stage directions — pure spoken script
- Never reproduce source content — fully original ideas
- Start immediately with the hook

Output ONLY the script text.""",
            },
        ]
        script = self._call_api(messages, temperature=0.7, max_tokens=8192).strip()
        word_count = len(script.split())
        return {
            "script": script,
            "word_count": word_count,
            "target_word_count": target_wc,
        }

    # ── State 6: Visual Sample Analysis ─────────────────────────────────

    def analyze_visual_samples(self, images_b64: List[str], mime_types: List[str] = None) -> Dict:
        if mime_types is None:
            mime_types = ["image/jpeg"] * len(images_b64)

        content: List[Dict] = [
            {
                "type": "text",
                "text": """Analyze these video frame screenshots and extract the channel's Visual Style Profile.

Output as JSON with EXACTLY these keys:
{
  "art_direction": "Overall aesthetic and visual style",
  "color_palette": "Dominant colors and tones",
  "lighting": "Type, direction, and quality of lighting",
  "camera_style": "Angle, distance, movement type",
  "composition": "Framing, symmetry, depth rules",
  "detail_density": "Visual complexity and detail level",
  "mood": "Overall atmosphere and emotional feel",
  "production_level": "Budget and quality indicators"
}

Output only JSON.""",
            }
        ]
        for b64, mime in zip(images_b64, mime_types):
            content.append(
                {
                    "type": "image_url",
                    "image_url": {"url": f"data:{mime};base64,{b64}"},
                }
            )

        result = self._call_api([{"role": "user", "content": content}], temperature=0.2)
        return self._parse_json(result)

    # ── State 7: Image Prompt Generation ────────────────────────────────

    def generate_image_prompts(self, script: str, visual_profile: Dict) -> List[Dict]:
        FERN_STYLE = (
            "dark cinematic 3D animation, featureless mannequin figure with no face no eyes no mouth, "
            "smooth blank head, faceless humanoid form, matte dark surface texture, "
            "dramatic chiaroscuro lighting, deep shadow, near-black color palette with selective accent color, "
            "cinematic wide or close-up shot, ominous atmosphere, rendered 3D scene, no text, no logos"
        )
        messages = [
            {
                "role": "system",
                "content": (
                    "You are an elite visual director specializing in dark 3D animation in the style of the YouTube channel Fern. "
                    "Every scene uses featureless mannequin figures with no facial traits, deep shadow lighting, and a cinematic 3D render aesthetic. "
                    "Segment scripts into 3-5 second beats and generate detailed standalone prompts. Output pure JSON only."
                ),
            },
            {
                "role": "user",
                "content": f"""Segment this script into beats of 3-5 seconds (8-12 words each) and generate one video prompt per beat.

SCRIPT:
{script}

MANDATORY VISUAL STYLE — every single prompt must include ALL of these:
{FERN_STYLE}

CHANNEL VISUAL PROFILE (additional style context):
{json.dumps(visual_profile, indent=2)}

Output a JSON array. Each element must have EXACTLY these keys:
{{
  "beat_number": 1,
  "segment": "exact script text for this beat",
  "image_prompt": "complete self-contained scene description — must include: featureless mannequin figure, dark 3D render, cinematic lighting, environment, camera framing, mood. NO realistic human faces ever.",
  "camera_angle": "specific angle and framing",
  "lighting": "type, direction, quality of shadow and light",
  "mood": "emotional atmosphere",
  "action": "what the mannequin figure or environment is doing",
  "visual_style": "dark 3D animation, faceless figure, cinematic render"
}}

CRITICAL: Every prompt must begin with 'dark cinematic 3D animation, featureless mannequin figure with no face,' — no exceptions. No realistic humans. No faces. Ever.

Output only the JSON array.""",
            },
        ]
        result = self._call_api(messages, temperature=0.6, max_tokens=16000)
        return self._parse_json(result)

    # ── State 8: Video Prompt Generation ────────────────────────────────

    def generate_video_prompts(self, image_prompts: List[Dict]) -> List[Dict]:
        compact = [
            {
                "beat_number": p["beat_number"],
                "image_prompt": p["image_prompt"],
                "mood": p.get("mood", ""),
                "action": p.get("action", ""),
            }
            for p in image_prompts
        ]
        messages = [
            {
                "role": "system",
                "content": "You are an elite video director. Generate motion prompts for image stills. Output pure JSON only.",
            },
            {
                "role": "user",
                "content": f"""For each beat, generate a video motion prompt describing movement, pacing, and camera motion.

BEATS:
{json.dumps(compact, indent=2)}

Output a JSON array with EXACTLY these keys per element:
{{
  "beat_number": 1,
  "video_prompt": "full motion description including camera movement, subject motion, timing, and transition style"
}}

Output only the JSON array.""",
            },
        ]
        result = self._call_api(messages, temperature=0.6, max_tokens=16000)
        additions = self._parse_json(result)

        video_map = {v["beat_number"]: v["video_prompt"] for v in additions}
        return [{**p, "video_prompt": video_map.get(p["beat_number"], "")} for p in image_prompts]

    # ── State 9: Thumbnail Analysis ──────────────────────────────────────

    def analyze_thumbnails(self, images_b64: List[str], mime_types: List[str] = None) -> Dict:
        if mime_types is None:
            mime_types = ["image/jpeg"] * len(images_b64)

        content: List[Dict] = [
            {
                "type": "text",
                "text": """Analyze these YouTube thumbnails and extract the channel's thumbnail design language.

Output as JSON with EXACTLY these keys:
{
  "typography_style": "font weight, case, size hierarchy",
  "layout": "compositional structure and arrangement",
  "color_contrast": "background approach and contrast strategy",
  "character_framing": "facial expression or character use",
  "emotional_trigger": "click psychology and psychological driver",
  "dominant_colors": ["color1", "color2", "color3"],
  "text_placement": "where and how text appears",
  "overall_style": "one-paragraph synthesis of thumbnail DNA"
}

Output only JSON.""",
            }
        ]
        for b64, mime in zip(images_b64, mime_types):
            content.append(
                {
                    "type": "image_url",
                    "image_url": {"url": f"data:{mime};base64,{b64}"},
                }
            )

        result = self._call_api([{"role": "user", "content": content}], temperature=0.2)
        return self._parse_json(result)

    # ── State 10: Thumbnail Concepts ─────────────────────────────────────

    def generate_thumbnail_concepts(
        self, thumbnail_dna: Dict, topic: str, script: str
    ) -> List[Dict]:
        preview = script[:500] + "..." if len(script) > 500 else script
        messages = [
            {
                "role": "system",
                "content": (
                    "You are an elite YouTube thumbnail strategist. "
                    "Generate high-CTR thumbnail concepts matched to channel DNA. Output pure JSON only."
                ),
            },
            {
                "role": "user",
                "content": f"""Generate 5 distinct thumbnail concepts for this video.

TOPIC: {topic}

THUMBNAIL DNA:
{json.dumps(thumbnail_dna, indent=2)}

SCRIPT OPENING:
{preview}

Output a JSON array of 5 objects, each with EXACTLY these keys:
{{
  "concept_number": 1,
  "visual_concept": "Detailed scene or composition description",
  "text_overlay": "Exact wording with placement and styling logic",
  "emotion_trigger": "The psychological driver behind the click",
  "generation_prompt": "Complete AI image generation prompt matching the channel's thumbnail DNA"
}}

Make each concept distinct in approach but all matched to the thumbnail DNA.
Output only the JSON array.""",
            },
        ]
        result = self._call_api(messages, temperature=0.8)
        return self._parse_json(result)

    # ── Content Calendar ─────────────────────────────────────────────────

    def generate_content_calendar(self, style_dna: Dict, n_videos: int = 5) -> List[Dict]:
        """
        Generate n fully original video concepts in the channel's style.
        Each concept includes title, hook, angle, topic, and thumbnail idea.
        """
        prompt = f"""You are a YouTube content strategist.

Based on the channel's Style DNA below, generate {n_videos} original video concepts.
These must be completely original ideas — not based on any specific existing video.
They should feel native to this channel's voice, niche, and audience.

Style DNA:
{json.dumps(style_dna, indent=2)}

Return a JSON array of exactly {n_videos} objects. Each object must have:
- "title": the video title written in this channel's naming style
- "topic": the core subject/theme in one phrase
- "hook": the opening 2–3 sentences that would grab the viewer instantly, in this channel's voice
- "angle": what makes this video unique vs generic content on the same topic
- "thumbnail_concept": a one-sentence description of the thumbnail visual

Return only the JSON array, no other text."""

        messages = [{"role": "user", "content": prompt}]
        result = self._call_api(messages, temperature=0.9, max_tokens=4096)
        return self._parse_json(result)

    # ── State 11: Word Document Export ───────────────────────────────────

    def export_to_docx(self, session_data: Dict) -> bytes:
        try:
            from docx import Document
            from docx.shared import Pt
            from docx.enum.text import WD_ALIGN_PARAGRAPH
        except ImportError:
            raise ImportError("python-docx is required for Word export: pip install python-docx")

        doc = Document()

        heading = doc.add_heading("YouTube Channel Clone Engine", 0)
        heading.alignment = WD_ALIGN_PARAGRAPH.CENTER

        doc.add_paragraph(f"Channel: {session_data.get('channel_url', 'N/A')}")
        doc.add_paragraph(f"Topic: {session_data.get('topic', 'N/A')}")
        doc.add_paragraph("")

        # Style DNA
        if session_data.get("style_dna"):
            doc.add_heading("Style DNA Analysis", 1)
            dna = session_data["style_dna"]
            table = doc.add_table(rows=1, cols=2)
            table.style = "Table Grid"
            hdr = table.rows[0].cells
            hdr[0].text = "Dimension"
            hdr[1].text = "Analysis"
            dimension_labels = [
                ("niche", "Niche"),
                ("target_audience", "Target Audience"),
                ("hook_architecture", "Hook Architecture"),
                ("script_flow", "Script Flow"),
                ("sentence_rhythm", "Sentence Rhythm"),
                ("tone_profile", "Tone Profile"),
                ("transition_techniques", "Transition Techniques"),
                ("curiosity_gaps", "Curiosity Gaps"),
                ("emotional_triggers", "Emotional Triggers"),
                ("retention_devices", "Retention Devices"),
                ("direct_address_style", "Direct Address Style"),
                ("words_per_second", "Words Per Second"),
                ("target_word_count", "Target Word Count"),
                ("overall_voice", "Overall Voice"),
            ]
            for key, label in dimension_labels:
                if key in dna:
                    row = table.add_row().cells
                    row[0].text = label
                    row[1].text = str(dna[key])
            doc.add_paragraph("")

        # Script
        if session_data.get("script"):
            doc.add_heading("Video Script", 1)
            wc = session_data.get("word_count", 0)
            twc = session_data.get("target_word_count", 0)
            doc.add_paragraph(f"Word Count: {wc} / {twc} target")
            doc.add_paragraph("")
            for para in session_data["script"].split("\n"):
                if para.strip():
                    doc.add_paragraph(para)
            doc.add_paragraph("")

        # Visual Profile
        if session_data.get("visual_profile"):
            doc.add_heading("Visual Style Profile", 1)
            for key, val in session_data["visual_profile"].items():
                p = doc.add_paragraph()
                run = p.add_run(f"{key.replace('_', ' ').title()}: ")
                run.bold = True
                p.add_run(str(val))
            doc.add_paragraph("")

        # Image Prompts
        if session_data.get("image_prompts"):
            doc.add_heading("Image Prompts Per Beat", 1)
            for beat in session_data["image_prompts"]:
                doc.add_heading(f"Beat {beat.get('beat_number', '?')}", 3)
                for field, label in [
                    ("segment", "Script Segment"),
                    ("image_prompt", "Image Prompt"),
                    ("camera_angle", "Camera Angle"),
                    ("lighting", "Lighting"),
                    ("mood", "Mood"),
                    ("action", "Action"),
                    ("visual_style", "Visual Style"),
                    ("video_prompt", "Video Prompt"),
                ]:
                    if beat.get(field):
                        p = doc.add_paragraph()
                        p.add_run(f"{label}: ").bold = True
                        p.add_run(beat[field])
                doc.add_paragraph("")

        # Thumbnail DNA
        if session_data.get("thumbnail_dna"):
            doc.add_heading("Thumbnail DNA", 1)
            for key, val in session_data["thumbnail_dna"].items():
                p = doc.add_paragraph()
                p.add_run(f"{key.replace('_', ' ').title()}: ").bold = True
                p.add_run(", ".join(val) if isinstance(val, list) else str(val))
            doc.add_paragraph("")

        # Thumbnail Concepts
        if session_data.get("thumbnail_concepts"):
            doc.add_heading("Thumbnail Concepts", 1)
            for concept in session_data["thumbnail_concepts"]:
                doc.add_heading(f"Concept {concept.get('concept_number', '?')}", 3)
                for field, label in [
                    ("visual_concept", "Visual Concept"),
                    ("text_overlay", "Text Overlay"),
                    ("emotion_trigger", "Emotion Trigger"),
                    ("generation_prompt", "Generation Prompt"),
                ]:
                    if concept.get(field):
                        p = doc.add_paragraph()
                        p.add_run(f"{label}: ").bold = True
                        p.add_run(concept[field])
                doc.add_paragraph("")

        buf = BytesIO()
        doc.save(buf)
        buf.seek(0)
        return buf.read()
